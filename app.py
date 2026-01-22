from flask import Flask, render_template, redirect, url_for, request, flash, jsonify
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
from werkzeug.middleware.proxy_fix import ProxyFix
import os
import re
import random
import uuid
import threading
import time
from datetime import datetime, timedelta, date
from functools import wraps
from collections import defaultdict

import requests
from bs4 import BeautifulSoup
from PIL import Image
try:
    import pytesseract
except ImportError:
    pytesseract = None
try:
    import docx
except ImportError:
    docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
from hijri_converter import Gregorian
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from config import Config
from models import db, User, Habit, HabitLog, Schedule, RoutineItem, ScheduleLog, PrayerLog, Dua, Day, IslamicEvent

app = Flask(__name__)
app.config.from_object(Config)

# Correctly handle proxy headers (Client -> Cloudflare -> Render -> App)
# We trust 2 proxies: Render and Cloudflare
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=2, x_proto=2, x_host=2, x_port=2, x_prefix=2)

db.init_app(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Helper for Local Time (GMT+6)
def get_local_now():
    """Returns the current time in GMT+6 (local time)."""
    # Using a fixed offset for simple implementation in this context
    return datetime.utcnow() + timedelta(hours=6)

def get_today():
    """Returns today's date in GMT+6."""
    return get_local_now().date()

def ensure_day(user_id, target_date):
    """
    Ensures a Day object exists for the given user and date.
    Returns the Day object.
    """
    day = Day.query.filter_by(user_id=user_id, date=target_date).first()
    if not day:
        try:
            day = Day(user_id=user_id, date=target_date)
            db.session.add(day)
            db.session.commit()
        except:
            db.session.rollback()
            day = Day.query.filter_by(user_id=user_id, date=target_date).first()
    return day

# Context Processor for current year/data
@app.route('/api/user/sync_local_ip', methods=['POST'])
@login_required
def sync_local_ip():
    data = request.get_json()
    local_ip = data.get('local_ip')
    fingerprint = data.get('fingerprint')
    
    if local_ip:
        current_user.local_ip = local_ip
    if fingerprint:
        current_user.device_fingerprint = fingerprint
        
    if local_ip or fingerprint:
        db.session.commit()
    return jsonify({"status": "success"})

@app.context_processor
def inject_now():
    return {
        'now': get_local_now(),
        'getattr': getattr
    }

def get_client_ip():
    """Returns the real client IP, prioritizing cloudflare and proxy headers."""
    # Check Cloudflare first as it's the most direct client IP source in this stack
    cf_ip = request.headers.get('Cf-Connecting-Ip') or request.headers.get('True-Client-Ip')
    if cf_ip:
        return cf_ip
    
    # Fallback to remote_addr (which is corrected by ProxyFix)
    return request.remote_addr

# --- Auth Routes ---
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password')
        
        if User.query.filter(db.func.lower(User.username) == username.lower()).first():
            flash('Username already exists', 'danger')
            return redirect(url_for('register'))
            
        if User.query.filter_by(email=email).first():
            flash('Email already registered. Please login or use a different email.', 'danger')
            return redirect(url_for('register'))
            
        user = User(username=username, email=email)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        login_user(user)
        return redirect(url_for('dashboard'))
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password')
        user = User.query.filter(db.func.lower(User.username) == username.lower()).first()
        
        if user:
            if user.is_banned:
                flash("Your account has been suspended by an administrator.", "danger")
                return redirect(url_for('login'))
            if user.check_password(password):
                login_user(user)
                return redirect(url_for('dashboard'))
            else:
                flash('Incorrect password. Please try again.', 'danger')
        else:
            flash('Username not found. Please check your spelling or register.', 'danger')
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/guest_login')
def guest_login():
    # Create a dummy guest user (NOT saved to DB)
    # Flask-Login requires a user object.
    # We will use a special 'GuestUser' class or just create a temporary DB user that gets cleaned up?
    # Better approach for prototype: Create a 'guest_x' account in DB automatically.
    import uuid
    guest_name = f"Guest_{uuid.uuid4().hex[:8]}"
    guest = User(username=guest_name, email=f"{guest_name}@temp.com", role='guest')
    guest.set_password('guest')
    db.session.add(guest)
    db.session.commit()
    login_user(guest)
    flash('Logged in as Guest. Data is temporary', 'info')
    return redirect(url_for('dashboard'))

@app.route('/api/user/location', methods=['POST'])
@login_required
def update_location_api():
    data = request.json
    lat = data.get('latitude')
    lon = data.get('longitude')
    
    if lat is not None and lon is not None:
        current_user.latitude = lat
        current_user.longitude = lon
        current_user.last_location_update = datetime.utcnow()
        db.session.commit()
        return jsonify({'success': True})
    return jsonify({'success': False, 'error': 'Missing coordinates'}), 400

# --- Admin Routes ---
# --- Admin Decorators ---

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            flash("Admin access required.", "danger")
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

def super_admin_required(f):
    # Alias for admin_required as we are removing super_admin role
    return admin_required(f)

@app.route('/admin')
@login_required
@admin_required
def admin_dashboard():
    stats = {
        'users': User.query.count(),
        'habits': Habit.query.count(),
        'prayers_logged': PrayerLog.query.count(),
        'guests': User.query.filter_by(role='guest').count()
    }
    users = User.query.order_by(User.join_date.desc()).limit(10).all()
    return render_template('admin_dashboard.html', stats=stats, users=users, active_tab='overview')

@app.route('/admin/users')
@login_required
@admin_required
def admin_users():
    search_query = request.args.get('q', '').strip()
    role_filter = request.args.get('role', '')
    status_filter = request.args.get('status', '')
    
    query = User.query
    
    if search_query:
        query = query.filter(
            (User.username.ilike(f'%{search_query}%')) | 
            (User.email.ilike(f'%{search_query}%')) |
            (db.cast(User.id, db.String).ilike(f'%{search_query}%'))
        )
    
    if role_filter:
        query = query.filter_by(role=role_filter)
        
    if status_filter:
        if status_filter == 'active':
            query = query.filter_by(is_banned=False)
        elif status_filter == 'banned':
            query = query.filter_by(is_banned=True)
            
    users = query.all()
    return render_template('admin_dashboard.html', users=users, active_tab='users', 
                           search_query=search_query, role_filter=role_filter, status_filter=status_filter)

@app.route('/admin/user/<int:user_id>/ban', methods=['POST'])
@login_required
@admin_required
def admin_ban_user(user_id):
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        flash("You cannot ban yourself!", "danger")
    else:
        user.is_banned = not user.is_banned
        db.session.commit()
        status = "banned" if user.is_banned else "unbanned"
        flash(f"User {user.username} has been {status}.", "success")
        # Audit log
        from models import AuditLog
        audit = AuditLog(admin_id=current_user.id, action='ban_user', target_user_id=user.id, 
                         reason=None, ip_address=get_client_ip(), local_ip=current_user.local_ip)
        db.session.add(audit)
        db.session.commit()
    return redirect(url_for('admin_users'))

@app.route('/admin/user/<int:user_id>/delete', methods=['POST'])
@login_required
@admin_required
def admin_delete_user(user_id):
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        flash("You cannot delete yourself!", "danger")
    else:
        # Audit log BEFORE deletion to keep record of who was deleted
        from models import AuditLog
        audit = AuditLog(admin_id=current_user.id, action='delete_user', target_user_id=user.id, 
                         reason="Admin manual deletion", ip_address=get_client_ip(), local_ip=current_user.local_ip)
        db.session.add(audit)
        
        db.session.delete(user)
        db.session.commit()
        flash(f"User {user.username} deleted.", "success")
    return redirect(url_for('admin_users'))

# --- Admin User Detail View ---
@app.route('/admin/user/<int:user_id>', methods=['GET'])
@login_required
@admin_required
def admin_user_detail(user_id):
    user = User.query.get_or_404(user_id)
    # Audit log for view action
    from models import AuditLog
    audit = AuditLog(admin_id=current_user.id, action='view_user', target_user_id=user.id, 
                     reason=None, ip_address=get_client_ip(), local_ip=current_user.local_ip)
    db.session.add(audit)
    db.session.commit()
    # Gather related data
    habits = Habit.query.filter_by(user_id=user.id).all()
    habit_logs = HabitLog.query.join(Habit).filter(Habit.user_id == user.id).all()
    prayer_logs = PrayerLog.query.filter_by(user_id=user.id).all()
    schedule_logs = ScheduleLog.query.filter_by(user_id=user.id).all()
    days = Day.query.filter_by(user_id=user.id).order_by(Day.date.desc()).limit(100).all()
    
    # [NEW] Fetch User's Active Schedule
    from models import Schedule
    active_schedule = Schedule.query.filter_by(user_id=user.id, is_active=True).first()
    
    return render_template('admin_user_detail.html', user=user, habits=habits, habit_logs=habit_logs,
                           prayer_logs=prayer_logs, schedule_logs=schedule_logs, days=days, 
                           active_schedule=active_schedule, active_tab='users')

@app.route('/admin/content/duas')
@login_required
@admin_required
def admin_duas():
    duas = Dua.query.filter_by(user_id=None).all()
    return render_template('admin_dashboard.html', duas=duas, active_tab='content_duas')

@app.route('/admin/content/events')
@login_required
@admin_required
def admin_events():
    events = IslamicEvent.query.all()
    return render_template('admin_dashboard.html', events=events, active_tab='content_events')

@app.route('/admin/content/duas/add', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_add_dua():
    if request.method == 'POST':
        title = request.form.get('title')
        category = request.form.get('category')
        arabic_text = request.form.get('arabic_text')
        english_meaning = request.form.get('english_meaning')
        
        new_dua = Dua(
            title=title,
            category=category,
            arabic_text=arabic_text,
            english_meaning=english_meaning,
            user_id=None # Global dua
        )
        db.session.add(new_dua)
        
        # Audit log
        from models import AuditLog
        audit = AuditLog(admin_id=current_user.id, action='add_dua', target_user_id=None, 
                         reason=f"Added global dua: {title}", ip_address=get_client_ip(), local_ip=current_user.local_ip)
        db.session.add(audit)
        
        db.session.commit()
        flash(f"Dua '{title}' added successfully!", "success")
        return redirect(url_for('admin_duas'))
    
    return render_template('admin_add_dua.html', active_tab='content')

@app.route('/admin/content/events/add', methods=['GET', 'POST'])
@login_required
@admin_required
def admin_add_event():
    if request.method == 'POST':
        title = request.form.get('title')
        date_str = request.form.get('date')
        color = request.form.get('color')
        
        event_date = datetime.strptime(date_str, '%Y-%m-%d')
        new_event = IslamicEvent(
            title=title,
            date=event_date,
            color=color
        )
        db.session.add(new_event)
        
        # Audit log
        from models import AuditLog
        audit = AuditLog(admin_id=current_user.id, action='add_event', target_user_id=None, 
                         reason=f"Added global event: {title}", ip_address=get_client_ip(), local_ip=current_user.local_ip)
        db.session.add(audit)
        
        db.session.commit()
        flash(f"Event '{title}' added to global calendar!", "success")
        return redirect(url_for('admin_events'))
        
    return render_template('admin_add_event.html', active_tab='content')

@app.route('/admin/content/duas/delete/<int:dua_id>', methods=['POST'])
@login_required
@admin_required
def admin_delete_dua(dua_id):
    dua = Dua.query.get_or_404(dua_id)
    title = dua.title
    
    # Audit log
    from models import AuditLog
    audit = AuditLog(admin_id=current_user.id, action='delete_dua', target_user_id=None, 
                     reason=f"Deleted global dua: {title}", ip_address=get_client_ip(), local_ip=current_user.local_ip)
    db.session.add(audit)
    
    db.session.delete(dua)
    db.session.commit()
    flash("Dua deleted.", "success")
    return redirect(url_for('admin_duas'))

@app.route('/admin/content/events/delete/<int:event_id>', methods=['POST'])
@login_required
@admin_required
def admin_delete_event(event_id):
    event = IslamicEvent.query.get_or_404(event_id)
    title = event.title
    
    # Audit log
    from models import AuditLog
    audit = AuditLog(admin_id=current_user.id, action='delete_event', target_user_id=None, 
                     reason=f"Deleted global event: {title}", ip_address=get_client_ip(), local_ip=current_user.local_ip)
    db.session.add(audit)
    
    db.session.delete(event)
    db.session.commit()
    flash("Event deleted.", "success")
    return redirect(url_for('admin_events'))

@app.route('/admin/logs')
@login_required
@admin_required
def admin_logs():
    from models import AuditLog
    logs = AuditLog.query.order_by(AuditLog.timestamp.desc()).limit(100).all()
    return render_template('admin_dashboard.html', logs=logs, active_tab='system_logs')

@app.route('/admin/log_action', methods=['POST'])
@login_required
@admin_required
def admin_log_action():
    data = request.get_json()
    action = data.get('action')
    target_user_id = data.get('target_user_id')
    reason = data.get('reason')
    
    from models import AuditLog
    audit = AuditLog(
        admin_id=current_user.id,
        action=action,
        target_user_id=target_user_id,
        reason=reason,
        ip_address=get_client_ip(),
        local_ip=current_user.local_ip
    )
    db.session.add(audit)
    db.session.commit()
    return jsonify({"status": "success"})

@app.route('/admin/system/cleanup', methods=['POST'])
@login_required
@admin_required
def admin_cleanup_guests():
    # Cleanup guests older than 24 hours
    cutoff = datetime.utcnow() - timedelta(hours=24)
    guests = User.query.filter(User.role == 'guest', User.join_date < cutoff).all()
    count = len(guests)
    for g in guests:
        db.session.delete(g)
    db.session.commit()
    flash(f"Cleared {count} temporary guest accounts.", "success")
    return redirect(url_for('admin_dashboard'))

# --- Main Routes ---
@app.route('/ping')
def ping():
    return "PONG", 200

def keep_alive():
    """
    Pings the app to prevent Render free instance from sleeping.
    """
    # Wait for app to boot
    time.sleep(10)
    url = os.environ.get('RENDER_EXTERNAL_URL')
    if not url:
        print("Keep-alive: RENDER_EXTERNAL_URL not found, skipping internal ping.")
        return
        
    print(f"Keep-alive: Monitoring {url}")
    while True:
        try:
            requests.get(f"{url}/ping")
            print("Keep-alive: Ping successful")
        except Exception as e:
            print(f"Keep-alive: Ping failed: {e}")
        time.sleep(600) # Ping every 10 minutes

# Start the keep-alive thread
if os.environ.get('RENDER'):
    threading.Thread(target=keep_alive, daemon=True).start()

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

def recalculate_day_score(day_id):
    """
    Sum up all points for a given day (Habits, Prayers, Schedule).
    """
    from models import HabitLog, PrayerLog, ScheduleLog, Day
    day = Day.query.get(day_id)
    if not day: return 0
    
    habit_points = db.session.query(db.func.sum(HabitLog.points)).filter(HabitLog.day_id == day_id).scalar() or 0
    prayer_points = db.session.query(db.func.sum(PrayerLog.spiritual_score)).filter(PrayerLog.day_id == day_id).scalar() or 0
    schedule_points = db.session.query(db.func.sum(ScheduleLog.points)).filter(ScheduleLog.day_id == day_id, ScheduleLog.status == True).scalar() or 0
    
    day.total_score = int(habit_points + prayer_points + schedule_points)
    db.session.commit()
    return day.total_score

@app.route('/dashboard', methods=['GET', 'POST'])
@login_required
def dashboard():
    today = get_today()
    current_day = ensure_day(current_user.id, today)
        
    # Handle Day Updates (Intention, Energy, Mood)
    if request.method == 'POST':
        # Check if it's a day update
        if 'intention' in request.form:
            current_day.intention = request.form.get('intention')
            current_day.energy_level = request.form.get('energy_level', type=int)
            current_day.mood = request.form.get('mood', type=int)
            current_day.reflection = request.form.get('reflection')
            db.session.commit()
            flash('Day updated!', 'success')
            return redirect(url_for('dashboard'))

    # Get habits for today
    habits = Habit.query.filter_by(user_id=current_user.id, is_paused=False).all()
    habit_logs = {}
    for h in habits:
        log = HabitLog.query.filter_by(habit_id=h.id, date=today).first()
        # Ensure log is linked to day
        if log and log.day_id is None:
            log.day_id = current_day.id
            db.session.commit()
            
        habit_logs[h.id] = log

    # Get Prayer status
    prayer_log = PrayerLog.query.filter_by(user_id=current_user.id, date=today).first()
    if not prayer_log:
        try:
            prayer_log = PrayerLog(user_id=current_user.id, date=today, day_id=current_day.id)
            db.session.add(prayer_log)
            db.session.commit()
        except:
            db.session.rollback()
            prayer_log = PrayerLog.query.filter_by(user_id=current_user.id, date=today).first()
    elif prayer_log.day_id is None:
        prayer_log.day_id = current_day.id
        db.session.commit()

    # Get Today's Schedule
    day_name = today.strftime('%A')
    # Assuming one active schedule for simplicity
    active_schedule = Schedule.query.filter_by(user_id=current_user.id, is_active=True).first()
    todays_routines = []
    schedule_status = {}
    if active_schedule:
        todays_routines = RoutineItem.query.filter_by(schedule_id=active_schedule.id, day_of_week=day_name).order_by(RoutineItem.start_time).all()
        for r in todays_routines:
            s_log = ScheduleLog.query.filter_by(routine_id=r.id, date=today).first()
            if s_log and s_log.day_id is None:
                s_log.day_id = current_day.id
                db.session.commit()
                
            schedule_status[r.id] = s_log.status if s_log else False

    return render_template('dashboard.html', 
                           habits=habits, 
                           habit_logs=habit_logs, 
                           prayer_log=prayer_log,
                           routines=todays_routines,
                           schedule_status=schedule_status,
                           today=today,
                           day=current_day)

@app.route('/api/day/update', methods=['POST'])
@login_required
def update_day_api():
    today = get_today()
    current_day = ensure_day(current_user.id, today)
    
    data = request.json
    if 'intention' in data:
        current_day.intention = data.get('intention')
    if 'energy_level' in data:
        current_day.energy_level = data.get('energy_level')
    if 'mood' in data:
        current_day.mood = data.get('mood')
    if 'reflection' in data:
        current_day.reflection = data.get('reflection')
        
    db.session.commit()
    return jsonify({'success': True, 'score': current_day.total_score})

# --- Habit Routes ---
@app.route('/habits')
@login_required
def habits_list():
    habits = Habit.query.filter_by(user_id=current_user.id).all()
    return render_template('add_habit.html', habits=habits) # Reusing add_habit for list/add for simplicity

@app.route('/habit/add', methods=['POST'])
@login_required
def add_habit():
    name = request.form.get('name')
    category = request.form.get('category')
    frequency = request.form.get('frequency')
    
    # V2 Fields
    target_value = request.form.get('target_value', 1, type=int)
    min_value = request.form.get('min_value', 1, type=int)
    priority = request.form.get('priority', 3, type=int)
    difficulty = request.form.get('difficulty', 1, type=int)
    identity_label = request.form.get('identity_label')
    
    # Calculate base points based on difficulty and priority
    # Formula: Base = 10 * Difficulty * Priority
    # This rewards harder, more important tasks.
    points = 10 * difficulty * ((priority + 1) / 2) # priority 1-5, multiplier 1-3x approx

    new_habit = Habit(
        name=name, 
        category=category, 
        frequency=frequency, 
        user_id=current_user.id,
        target_value=target_value,
        min_value=min_value,
        priority=priority,
        difficulty=difficulty,
        identity_label=identity_label,
        points=int(points)
    )
    db.session.add(new_habit)
    db.session.commit()
    flash('Habit created!', 'success')
    return redirect(url_for('habits_list'))

@app.route('/habit/edit/<int:habit_id>', methods=['GET', 'POST'])
@login_required
def edit_habit(habit_id):
    habit = Habit.query.get_or_404(habit_id)
    if habit.owner != current_user:
        flash('Unauthorized', 'danger')
        return redirect(url_for('dashboard'))
        
    if request.method == 'POST':
        habit.name = request.form.get('name')
        habit.category = request.form.get('category')
        habit.frequency = request.form.get('frequency')
        habit.priority = request.form.get('priority', type=int)
        habit.difficulty = request.form.get('difficulty', type=int)
        habit.target_value = request.form.get('target_value', type=int)
        habit.min_value = request.form.get('min_value', type=int)
        habit.unit = request.form.get('unit')
        habit.identity_label = request.form.get('identity_label')
        
        # Recalculate points
        habit.points = int(10 * habit.difficulty * ((habit.priority + 1) / 2))
        
        db.session.commit()
        flash('Habit updated!', 'success')
        return redirect(url_for('habits_list'))
        
    return render_template('edit_habit.html', habit=habit)

@app.route('/habit/delete/<int:habit_id>', methods=['POST'])
@login_required
def delete_habit(habit_id):
    habit = Habit.query.get_or_404(habit_id)
    if habit.owner != current_user:
        flash('Unauthorized', 'danger')
        return redirect(url_for('dashboard'))
        
    db.session.delete(habit)
    db.session.commit()
    flash('Habit deleted.', 'success')
    return redirect(url_for('habits_list'))

@app.route('/habit/toggle/<int:habit_id>', methods=['POST'])
@login_required
def toggle_habit(habit_id):
    habit = Habit.query.get_or_404(habit_id)
    if habit.owner != current_user:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
        
    today = get_today()
    current_day = ensure_day(current_user.id, today)

    log = HabitLog.query.filter_by(habit_id=habit.id, date=today).first()
    
    # Check if this is a simple toggle or increment
    is_multistep = habit.target_value > 1
    
    if log:
        if is_multistep:
            # Increment logic
            # If already at target, reset to 0 (Looping behavior for easy correction)
            # OR allow explicit undo. Let's do: Increment until target, then one more click resets to 0.
            if log.value_done >= habit.target_value:
                log.value_done = 0
                log.status = False
                log.points = 0
            else:
                log.value_done += 1
                # Check if now complete
                if log.value_done >= habit.target_value:
                    log.status = True
                    log.points = habit.points
                else:
                    log.status = False
                    # Pro-rated points? (TotalPoints * (Current / Target))
                    # Let's keep points simple: Only on completion? 
                    # User asked for "points for repeating".
                    # So we calculate fractional points.
                    points_per_unit = habit.points / habit.target_value
                    log.points = int(points_per_unit * log.value_done)
        else:
            # Simple Toggle (Binary)
            if log.status:
                log.status = False
                log.value_done = 0
                log.points = 0
            else:
                log.status = True
                log.value_done = habit.target_value
                log.points = habit.points

        if log.day_id is None: log.day_id = current_day.id
            
    else:
        # First interaction
        # If multistep, start at 1. Else start at target.
        initial_value = 1 if is_multistep else habit.target_value
        initial_status = (initial_value >= habit.target_value)
        
        # Points
        points = 0
        if initial_status:
            points = habit.points
        elif is_multistep:
             points = int((habit.points / habit.target_value) * 1)

        log = HabitLog(
            habit_id=habit.id, 
            date=today, 
            status=initial_status, 
            day_id=current_day.id,
            value_done=initial_value,
            points=points
        )
        db.session.add(log)
    
    db.session.commit()
    recalculate_day_score(current_day.id)
    
    return jsonify({
        'success': True, 
        'new_status': log.status, 
        'value_done': log.value_done, 
        'target_value': habit.target_value
    })

# --- Schedule Routes ---
@app.route('/schedule', methods=['GET', 'POST'])
@login_required
def schedule_view():
    active_schedule = Schedule.query.filter_by(user_id=current_user.id, is_active=True).first()
    if request.method == 'POST':
        # Create new schedule if none exists
        if not active_schedule:
            name = request.form.get('schedule_name', 'My Schedule')
            active_schedule = Schedule(name=name, user_id=current_user.id, is_active=True)
            db.session.add(active_schedule)
            db.session.commit()
            flash('Schedule created! Start adding your classes.', 'success')
            return redirect(url_for('schedule_view'))
            
        # Add Routine Item
        title = request.form.get('title')
        day = request.form.get('day')
        start = request.form.get('start_time')
        end = request.form.get('end_time')
        location = request.form.get('location')
        
        if not title or not start or not end:
             flash('Please fill in all required fields.', 'warning')
             return redirect(url_for('schedule_view'))
             
        start_time = datetime.strptime(start, '%H:%M').time()
        end_time = datetime.strptime(end, '%H:%M').time()
        
        item = RoutineItem(schedule_id=active_schedule.id, title=title, day_of_week=day, start_time=start_time, end_time=end_time, location=location)
        db.session.add(item)
        db.session.commit()
        flash('Class/Routine added!', 'success')
        return redirect(url_for('schedule_view'))

    return render_template('schedules.html', schedule=active_schedule)

@app.route('/schedule/toggle/<int:id>', methods=['POST'])
@login_required
def toggle_routine(id):
    item = RoutineItem.query.get_or_404(id)
    # Verify ownership via schedule
    if item.schedule.owner != current_user:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 403
        
    today = get_today()
    current_day = ensure_day(current_user.id, today)

    log = ScheduleLog.query.filter_by(routine_id=item.id, date=today).first()
    
    if log:
        log.status = not log.status
        if log.day_id is None: log.day_id = current_day.id
    else:
        log = ScheduleLog(routine_id=item.id, user_id=current_user.id, date=today, status=True, day_id=current_day.id)
        db.session.add(log)
    
    db.session.commit()
    recalculate_day_score(current_day.id)
    return jsonify({'success': True, 'new_status': log.status})

@app.route('/schedule/delete/<int:id>', methods=['POST'])
@login_required
def delete_routine(id):
    item = RoutineItem.query.get_or_404(id)
    if item.schedule.owner != current_user:
        flash('Unauthorized', 'danger')
        return redirect(url_for('schedule_view'))
        
    db.session.delete(item)
    db.session.commit()
    flash('Class/Routine deleted.', 'success')
    return redirect(url_for('schedule_view'))

@app.route('/schedule/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_routine(id):
    item = RoutineItem.query.get_or_404(id)
    if item.schedule.owner != current_user:
        flash('Unauthorized', 'danger')
        return redirect(url_for('schedule_view'))
        
    if request.method == 'POST':
        item.title = request.form.get('title')
        item.day_of_week = request.form.get('day')
        start = request.form.get('start_time')
        end = request.form.get('end_time')
        item.location = request.form.get('location')
        
        if start:
            item.start_time = datetime.strptime(start, '%H:%M').time()
        if end:
            item.end_time = datetime.strptime(end, '%H:%M').time()
            
        db.session.commit()
        flash('Routine updated!', 'success')
        return redirect(url_for('schedule_view'))
        
    return render_template('edit_routine.html', item=item)

@app.route('/schedule/destroy/<int:id>', methods=['POST'])
@login_required
def delete_schedule(id):
    schedule = Schedule.query.get_or_404(id)
    if schedule.owner != current_user:
        flash('Unauthorized', 'danger')
        return redirect(url_for('schedule_view'))
        
    db.session.delete(schedule)
    db.session.commit()
    flash(f'Schedule "{schedule.name}" deleted successfully.', 'success')
    return redirect(url_for('schedule_view'))

# --- Islamic/Prayer Routes ---
@app.route('/prayers', methods=['GET', 'POST'])
@login_required
def prayers():
    today = get_today()
    current_day = ensure_day(current_user.id, today)

    log = PrayerLog.query.filter_by(user_id=current_user.id, date=today).first()
    if not log:
        try:
            log = PrayerLog(user_id=current_user.id, date=today, day_id=current_day.id)
            db.session.add(log)
            db.session.commit()
        except:
            db.session.rollback()
            log = PrayerLog.query.filter_by(user_id=current_user.id, date=today).first()
    elif log.day_id is None:
        log.day_id = current_day.id
        db.session.commit()
        
    if request.method == 'POST':
        data = request.json
        prayer_name = data.get('prayer') # fajr, dhuhr...
        status = data.get('status') # boolean
        
        if hasattr(log, prayer_name):
            setattr(log, prayer_name, status)
            
            # Recalculate Score (Basic Logic)
            score = 0
            for p in ['fajr', 'dhuhr', 'asr', 'maghrib', 'isha']:
                if getattr(log, p):
                    score += 100 # 100 * 5 = 500 base
            log.spiritual_score = score
            
            db.session.commit()
            recalculate_day_score(log.day_id)
            return jsonify({'success': True, 'score': score})
            
    return render_template('prayers.html', log=log)

@app.route('/islamic')
@login_required
def islamic_hub():
    # Show Duas, Calendar link, etc.
    duas = Dua.query.filter((Dua.user_id == None) | (Dua.user_id == current_user.id)).all()
    return render_template('dua.html', duas=duas)

@app.route('/init_db')
def init_db():
    db.create_all()
    
    # Seed Duas if empty
    # Seed Duas if not enough
    if Dua.query.count() < 3:
        seed_duas = [
            # Dua(title="Morning Dua", ...), # Already exists likely
            # ... add the new ones ...
        ]
        # To avoid duplicates, let's just use a smarter list of "all intended duas" and check existence.
        
        candidates = [
            Dua(title="Morning Dua", arabic_text="ٱلْحَمْدُ لِلَّٰهِ ٱلَّذِي أَحْيَانَا بَعْدَ مَا أَمَاتَنَا وَإِلَيْهِ ٱلنُّشُورُ", english_meaning="All praise is for Allah who gave us life after having taken it from us and unto Him is the resurrection.", bangla_meaning="সমস্ত প্রশংসা আল্লাহর জন্য, যিনি আমাদের মৃত্যুর পর জীবন দান করেছেন এবং তাঁরই দিকে আমাদের পুনরুত্থান।", category="Morning"),
            Dua(title="Before Sleep", arabic_text="بِاسْمِكَ رَبِّ وَضَعْتُ جَنْبِي، وَبِكَ أَرْفَعُهُ", english_meaning="In Your name my Lord, I lie down, and in Your name I rise.", bangla_meaning="হে আমার রব! আপনার নামেই আমি আমি আমার পার্শ্বদেশ বিছানায় রাখলাম এবং আপনার নামেই আমি তা উঠাবো।", category="Evening"),
            Dua(title="For Knowledge", arabic_text="رَّبِّ زِدْنِى عِلْمًا", english_meaning="My Lord, increase me in knowledge.", bangla_meaning="হে আমার প্রতিপালক! আমার জ্ঞান বৃদ্ধি করে দিন।", category="Knowledge"),
            Dua(title="For Parents", arabic_text="رَّبِّ ارْحَمْهُمَا كَمَا رَبَّيَانِي صَغِيرًا", english_meaning="My Lord, have mercy upon them [my parents] as they brought me up [when I was] small.", bangla_meaning="হে আমার প্রতিপালক! তাদের উভয়ের প্রতি দয়া করুন, যেমন তারা আমাকে শৈশবে লালন-পালন করেছেন।", category="Family"),
            Dua(title="Before Eating", arabic_text="بِسْمِ اللَّهِ", english_meaning="In the name of Allah.", bangla_meaning="আল্লাহর নামে।", category="Daily"),
            Dua(title="After Eating", arabic_text="الْحَمْدُ لِلَّهِ الَّذِي أَطْعَمَنَا وَسَقَانَا وَجَعَلَنَا مُسْلِمِينَ", english_meaning="All praise is due to Allah who fed us, gave us drink, and made us Muslims.", bangla_meaning="সকল প্রশংসা আল্লাহর জন্য যিনি আমাদের আহার করিয়েছেন, পান করিয়েছেন এবং মুসলিম বানিয়েছেন।", category="Daily"),
            Dua(title="Leaving Home", arabic_text="بِسْمِ اللهِ تَوَكَّلْتُ عَلَى اللهِ، لَا حَوْلَ وَلَا قُوَّةَ إِلَّا بِاللهِ", english_meaning="In the name of Allah, I place my trust in Allah; there is no might and no power except by Allah.", bangla_meaning="আল্লাহর নামে, আমি আল্লাহর উপর ভরসা করলাম। আল্লাহর সাহায্য ছাড়া কোন ক্ষমতা ও শক্তি নেই।", category="Travel"),
            Dua(title="Entering Mosque", arabic_text="اللَّهُمَّ افْتَحْ لِي أَبْوَابَ رَحْمَتِكَ", english_meaning="O Allah, open for me the doors of Your mercy.", bangla_meaning="হে আল্লাহ! আমার জন্য আপনার রহমতের দরজাগুলো খুলে দিন।", category="Mosque"),
            Dua(title="For Forgiveness", arabic_text="أَسْتَغْفِرُ اللَّهَ وَأَتُوبُ إِلَيْهِ", english_meaning="I seek forgiveness from Allah and turn to Him in repentance.", bangla_meaning="আমি আল্লাহর কাছে ক্ষমা প্রার্থনা করছি এবং তাঁর দিকে তওবা করছি।", category="Forgiveness"),
            Dua(title="When in Distress", arabic_text="يَا حَيُّ يَا قَيُّومُ بِرَحْمَتِكَ أَسْتَغِيثُ", english_meaning="O Ever Living, O Self-Subsisting and Supporter of all, by Your mercy I seek assistance.", bangla_meaning="হে চিরঞ্জীব, হে চিরস্থায়ী! আমি আপনার রহহমতের উসিলায় সাহায্য প্রার্থনা করছি।", category="General"),
            Dua(title="For Protection", arabic_text="بِسْمِ اللَّهِ الَّذِي لَا يَضُرُّ مَعَ اسْمِهِ شَيْءٌ فِي الْأَرْضِ وَلَا فِي السَّمَاءِ وَهُوَ السَّمِيعُ الْعَلِيمُ", english_meaning="In the name of Allah, with whose name nothing on earth or in the sky can harm. He is the All-Hearing, All-Knowing.", bangla_meaning="আল্লাহর নামে, যাঁর নামের বরকতে আসমান ও যমিনের কোনো কিছুই ক্ষতি করতে পারে না, তিনি সর্বশ্রোতা ও সর্বজ্ঞ।", category="Protection"),
            Dua(title="For Patience", arabic_text="رَبَّنَا أَفْرِغْ عَلَيْنَا صَبْرًا وَثَبِّتْ أَقْدَامَنَا وَانصُرْنَا عَلَى الْقَوْمِ الْكَافِرِينَ", english_meaning="Our Lord, pour upon us patience and plant firmly our feet and give us victory over the disbelieving people.", bangla_meaning="হে আমাদের প্রতিপালক! আমাদের ধৈর্য দান করুন, আমাদের পদযুগল অবিচলিত রাখুন এবং কাফের সম্প্রদায়ের বিরুদ্ধে আমাদের সাহায্য করুন।", category="Hardship"),
            Dua(title="For Ease", arabic_text="اللَّهُمَّ لَا سَهْلَ إِلَّا مَا جَعَلْتَهُ سَهْلًا، وَأَنْتَ تَجْعَلُ الْحَزْنَ إِذَا شِئْتَ سَهْلًا", english_meaning="O Allah, there is no ease except in what You have made easy, and You make the difficulty, if You will, easy.", bangla_meaning="হে আল্লাহ! আপনি যা সহজ করেছেন তা ছাড়া কোনো কিছুই সহজ নয়। আর যখন আপনি চান, তখন কঠিনকেও সহজ করে দেন।", category="Hardship"),
            Dua(title="For Anxiety", arabic_text="اللَّهُمَّ إِنِّي أَعُوذُ بِكَ مِنَ الْهَمِّ وَالْحَزَنِ، وَالْعَجْزِ وَالْكَسَلِ، وَالْبُخْلِ وَالْجُبْنِ، وَضَلَعِ الدَّيْنِ، وَغَلَبَةِ الرِّجَالِ", english_meaning="O Allah, I seek refuge in You from anxiety and sorrow, weakness and laziness, miserliness and cowardice, the burden of debts and from being overpowered by men.", bangla_meaning="হে আল্লাহ! আমি আপনার কাছে আশ্রয় চাই দুশ্চিন্তা ও দুঃখ থেকে, অক্ষমতা ও অলসতা থেকে, কৃপণতা ও ভীরুতা থেকে, ঋণের বোঝা ও মানুষের প্রাধান্য বিস্তার থেকে।", category="Anxiety"),
            Dua(title="Sayyidul Istighfar", arabic_text="اللَّهُمَّ أَنْتَ رَبِّي لَا إِلَهَ إِلَّا أَنْتَ، خَلَقْتَنِي وَأَنَا عَبْدُكَ، وَأَنَا عَلَى عَهْدِكَ وَوَعْدِكَ مَا اسْتَطَعْتُ، أَعُوذُ بِكَ مِنْ شَرِّ مَا صَنَعْتُ، أَبُوءُ لَكَ بِنِعْمَتِكَ عَلَيَّ، وَأَبُوءُ لَكَ بِذَنْبِي فَاغْفِرْ لِي فَإِنَّهُ لَا يَغْفِرُ الذُّنُوبَ إِلَّا أَنْتَ", english_meaning="O Allah, You are my Lord, none has the right to be worshipped except You. You created me and I am Your servant, and I abide by Your covenant and promise as best I can. I seek refuge in You from the evil, which I have committed. I acknowledge Your blessing upon me and I acknowledge my sin, so forgive me, for none can forgive sins except You.", bangla_meaning="হে আল্লাহ! আপনি আমার প্রতিপালক। আপনি ছাড়া আর কোনো উপাস্য নেই। আপনিই আমাকে সৃষ্টি করেছেন এবং আমি আপনার বান্দা। আমি আমার সাধ্যমতো আপনার সাথে কৃত অঙ্গীকার ও প্রতিশ্রুতির ওপর কায়েম আছি। আমার কৃতকর্মের অনিষ্ট থেকে আমি আপনার কাছে আশ্রয় চাই। আপনার যে নিয়ামত আমার ওপর রয়েছে আমি তা স্বীকার করছি এবং আমি আমার গুনাহও স্বীকার করছি। সুতরাং আপনি আমাকে ক্ষমা করে দিন। কেননা আপনি ছাড়া আর কেউ গুনাহ ক্ষমা করতে পারে না।", category="Forgiveness"),
            Dua(title="Ayatul Kursi", arabic_text="اللَّهُ لَا إِلَهَ إِلَّا هُوَ الْحَيُّ الْقَيُّومُ ۚ لَا تَأْخُذُهُ سِنَةٌ وَلَا نَوْمٌ ۚ لَهُ مَا فِي السَّمَاوَاتِ وَمَا فِي الْأَرْضِ ۗ مَنْ ذَا الَّذِي يَشْفَعُ عِنْدَهُ إِلَّا بِإِذْنِهِ ۚ يَعْلَمُ مَا بَيْنَ أَيْدِيهِمْ وَمَا خَلْفَهُمْ ۖ وَلَا يُحِيطُونَ بِشَيْءٍ مِنْ عِلْمِهِ إِلَّا بِمَا شَاءَ ۚ وَسِعَ كُرْسِيُّهُ السَّمَاوَاتِ وَالْأَرْضَ ۖ وَلَا يَئُودُهُ حِفْظُهُمَا ۚ وَهُوَ الْعَلِيُّ الْعَظِيمُ", english_meaning="Allah - there is no deity except Him, the Ever-Living, the Sustainer of [all] existence. Neither drowsiness overtakes Him nor sleep. To Him belongs whatever is in the heavens and whatever is on the earth. Who is it that can intercede with Him except by His permission? He knows what is [presently] before them and what will be after them, and they encompass not a thing of His knowledge except for what He wills. His Kursi extends over the heavens and the earth, and their preservation tires Him not. And He is the Most High, the Most Great.", bangla_meaning="আল্লাহ ছাড়া অন্য কোনো উপাস্য নেই, তিনি জীবিত ও সবকিছুর ধারক। তাঁকে তন্দ্রা ও নিদ্রা স্পর্শ করতে পারে না। আসমান ও যমীনে যা কিছু আছে সবই তাঁর। কে আছে এমন যে তাঁর অনুমতি ছাড়া তাঁর কাছে সুপারিশ করবে? তাদের সামনে ও পিছনে যা কিছু আছে তা তিনি জানেন। তারা তাঁর জ্ঞানের সামান্য অংশও আয়ত্ত করতে পারে না, তবে তিনি যতটুকু চান তা ছাড়া। তাঁর কুরসী আসমান ও যমীনব্যাপী পরিব্যাপ্ত। এ দুটোর রক্ষণাবেক্ষণ তাঁকে ক্লান্ত করে না। তিনি সুউচ্চ ও মহান।", category="Protection")
        ]
        
        for d in candidates:
            exists = Dua.query.filter_by(title=d.title).first()
            if not exists:
                db.session.add(d)
        db.session.commit()
        


    # Create Default Admin if not exists
    admin = User.query.filter_by(username='admin').first()
    if not admin:
        admin = User(username='admin', email='admin@habit.local', role='admin')
        admin.set_password('adminpass')
        db.session.add(admin)
        
    db.session.commit()
    return "Database Initialized, Seeded & Admin Created!"

# Static Islamic Events Dictionary
ISLAMIC_EVENTS = {
    'Muharram': {
        1: 'Islamic New Year (1st Muharram) - Commemorates the Hijra',
        10: ['Ashura - The Saving of Prophet Musa (AS) from Pharaoh', 'Martyrdom of Imam Husayn (AS) at Karbala']
    },
    'Safar': {
        1: 'Khalid bin Walid (RA) embraced Islam (8 AH)',
        27: 'Migration of Prophet Muhammad (SAW) to Madinah (1 AH)'
    },
    'Rabi al-Awwal': {
        1: 'Hijrah: Prophet Muhammad (SAW) leaves Mecca for Medina',
        12: 'Mawlid al-Nabi (Sunni) - Birth of Prophet Muhammad (SAW)',
        17: 'Mawlid al-Nabi (Shia) - Birth of Prophet Muhammad (SAW)'
    },
    'Rabi al-Thani': {
        11: 'Urs of Abdul Qadir Gilani (Ghaus-e-Azam)'
    },
    'Jumada al-Ula': {
         5: 'Birth of Zainab bint Ali (AS)',
         15: 'Birth of Imam Zayn al-Abidin (AS)'
    },
    'Jumada al-Ahirah': {
        20: 'Birth of Fatima Zahra (AS)'
    },
    'Rajab': {
        1: 'Birth of Imam Muhammad al-Baqir (AS)', 
        13: 'Birth of Imam Ali ibn Abi Talib (AS) inside the Kaaba', 
        15: [
            'Change of the Qibla from Jerusalem to Mecca (2 AH)', 
            'Death Anniversary of Sayyida Zainab bint Ali (AS) (62 AH)',
            'Termination of the Boycott of Shi\'b Abi Talib (Social Sanctions ended)'
        ],
        27: 'Isra and Mi\'raj (The Night Journey and Ascension)'
    },
    'Shaban': {
        3: 'Birth of Imam Husayn (AS)',
        15: 'Mid-Shaban (Laylat al-Bara\'at) - Night of Forgiveness'
    },
    'Ramadan': {
        1: 'First Day of Ramadan', 
        10: 'Death of Khadija bint Khuwaylid (RA)',
        15: 'Birth of Imam Hasan (AS)',
        17: 'Battle of Badr (2 AH) - First decisive battle in Islam', 
        19: 'Injury of Imam Ali (AS) in Kufa',
        21: 'Martyrdom of Imam Ali (AS)', 
        27: 'Laylat al-Qadr (The Night of Decree) - Most probable date'
    },
    'Shawwal': {
        1: 'Eid al-Fitr - Festival of Breaking the Fast',
        8: 'Destruction of Baqi Cemetery (Yaum-e-Gham)'
    },
    'Dhu al-Qadah': {
        11: 'Birth of Imam Ali al-Rida (AS)'
    },
    'Dhu al-Hijjah': {
        1: 'Marriage of Ali (AS) and Fatima (AS)', 
        7: 'Martyrdom of Imam Muhammad al-Baqir (AS)',
        9: 'Day of Arafah - The pinnacle of Hajj', 
        10: 'Eid al-Adha - Festival of Sacrifice', 
        18: 'Eid al-Ghadir - Appointment of Imam Ali (AS)'
    }
}


# [START] Routes for Calendar, Upload, and APIs

@app.route('/calendar')
@login_required
def calendar_view():
    return render_template('new_calendar.html')

@app.route('/api/events')
@login_required
def get_calendar_events():
    events = []
    
    # FullCalendar sends start and end as ISO date strings
    start_str = request.args.get('start', '').split('T')[0]
    end_str = request.args.get('end', '').split('T')[0]
    
    if not start_str or not end_str:
        return jsonify([])

    try:
        start_date = datetime.strptime(start_str, '%Y-%m-%d').date()
        end_date = datetime.strptime(end_str, '%Y-%m-%d').date()
    except ValueError:
        return jsonify([])

    # 1. Schedule Items (Persistent until deleted)
    active_schedule = Schedule.query.filter_by(user_id=current_user.id, is_active=True).first()
    if active_schedule:
        days_map = {'Monday': 0, 'Tuesday': 1, 'Wednesday': 2, 'Thursday': 3, 'Friday': 4, 'Saturday': 5, 'Sunday': 6}
        routines = RoutineItem.query.filter_by(schedule_id=active_schedule.id).all()
        
        curr = start_date
        while curr <= end_date:
            day_name = curr.strftime('%A')
            for r in routines:
                if r.day_of_week == day_name:
                    events.append({
                        'title': r.title,
                        'start': f"{curr}T{r.start_time}",
                        'end': f"{curr}T{r.end_time}",
                        'color': '#6366f1' # Primary color
                    })
            curr += timedelta(days=1)
                    
    # 3. Islamic Special Days
    # Static fallback
    special_days = [
        {'title': 'Ramadan Begins', 'start': '2025-03-01', 'color': '#10b981'},
        {'title': 'Eid al-Fitr', 'start': '2025-03-30', 'color': '#f59e0b'},
    ]
    events.extend(special_days)
    
    # Dynamic from DB
    db_events = IslamicEvent.query.filter(IslamicEvent.date >= start_date, IslamicEvent.date <= end_date).all()
    for de in db_events:
        events.append({
            'title': de.title,
            'start': de.date.isoformat(),
            'color': de.color,
            'description': de.description
        })
    
    return jsonify(events)

@app.route('/api/day_details')
@login_required
def get_day_details():
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date required'}), 400
    
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        month_name = dt.strftime('%B').lower()
        day_num = dt.day
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        
        events = []
        
        # 0. Islamic Date and Dynamic Significance
        try:
            h_date = Gregorian(dt.year, dt.month, dt.day).to_hijri()
            h_month_name = h_date.month_name()
            h_day_str = h_day_suffix(h_date.day)
            
            islamic_date_str = f"{h_day_str} {h_month_name}"
            events.append(f"📅 Islamic Date: {islamic_date_str}, {h_date.year} AH")
            
            # Lookup in Static Events
            month_events = ISLAMIC_EVENTS.get(h_month_name, {})
            day_data = month_events.get(h_date.day)
            
            if day_data:
                if isinstance(day_data, list):
                    for event in day_data:
                        events.append(f"🕌 {event}")
                else:
                    events.append(f"🕌 {day_data}")
            
            # Dynamic Fallback: Wikipedia Scrape (Month Page)
            # Scrape the month page (e.g. wiki/Rajab) and look for the specific day
            try:
                # wiki/Rajab
                safe_month = h_month_name.replace(' ', '_')
                wiki_url = f"https://en.wikipedia.org/wiki/{safe_month}"
                
                resp = requests.get(wiki_url, headers=headers, timeout=3)
                if resp.status_code == 200:
                    soup = BeautifulSoup(resp.content, 'html.parser')
                    # Find all list items
                    # Look for day pattern: "15 Rajab" or "15th" or just "15 " at start
                    day_patterns = [
                        f"{h_date.day} {h_month_name}", 
                        f"{h_date.day}th {h_month_name}",
                        f"{h_date.day}st {h_month_name}",
                        f"{h_date.day}nd {h_month_name}",
                        f"{h_date.day}rd {h_month_name}",
                        f"{h_date.day} " # Simple day number check (risky but getting hit is better)
                    ]
                    
                    content_div = soup.find('div', {'id': 'mw-content-text'})
                    if content_div:
                        for li in content_div.find_all('li'):
                            text = li.get_text(strip=True)
                            # Check if line starts with day number
                            # e.g. "27 Rajab – Isra..." or "27 – Isra..."
                            if any(text.lower().startswith(p.lower()) for p in day_patterns) or text.startswith(f"{h_date.day} –") or text.startswith(f"{h_date.day} -"):
                                # Avoid huge paragraphs, just take the line
                                if len(text) > 10 and len(text) < 300:
                                     # Clean citations [1]
                                    import re
                                    clean_text = re.sub(r'\[\d+\]', '', text)
                                    events.append(f"📜 {clean_text}")
                
                # Try specific Date page (e.g. wiki/15_Rajab)
                date_page_url = f"https://en.wikipedia.org/wiki/{h_date.day}_{safe_month}"
                d_resp = requests.get(date_page_url, headers=headers, timeout=3)
                if d_resp.status_code == 200:
                    d_soup = BeautifulSoup(d_resp.content, 'html.parser')
                    for section_header in ['Events', 'Observances', 'Births', 'Deaths']:
                        s_h = d_soup.find('span', {'id': section_header})
                        if s_h:
                            ul = s_h.parent.find_next_sibling('ul') or s_h.parent.find_next('ul')
                            if ul:
                                for li in ul.find_all('li')[:3]:
                                    events.append(f"🏷️ {li.get_text(strip=True)}")
            except Exception as w_e:
                 # Silently fail
                 print(f"Wiki Search Error: {w_e}")

            # Fallback for events: WikiShia
            if len(events) <= 1: # Only header exists
                 try:
                     wikishia_url = f"https://en.wikishia.net/view/{h_month_name}"
                     w_resp = requests.get(wikishia_url, headers=headers, timeout=3)
                     if w_resp.status_code == 200:
                         w_soup = BeautifulSoup(w_resp.content, 'html.parser')
                         for li in w_soup.find_all('li'):
                             text = li.get_text(strip=True)
                             if text.startswith(f"{h_date.day} {h_month_name}") or text.startswith(f"{h_date.day}th {h_month_name}"):
                                 events.append(f"🕌 {text}")
                 except Exception as ws_e:
                     print(f"WikiShia Error: {ws_e}")

            if len(events) <= 1:
                events.append("No major historical events recorded for this date.")

            # Independent Daily Reflection (Quran/Hadith)
            hadith = None
            try:
                # 1. Fetch random Ayah with translation
                # Ayah number between 1 and 6236
                ayah_num = random.randint(1, 6236)
                q_res = requests.get(f'https://api.alquran.cloud/v1/ayah/{ayah_num}/editions/quran-uthmani,en.sahih', timeout=5)
                if q_res.ok:
                    q_data = q_res.json()
                    if q_data.get('data') and len(q_data['data']) >= 2:
                        arabic = q_data['data'][0]
                        english = q_data['data'][1]
                        
                        surah = arabic['surah']['englishName']
                        number = arabic['numberInSurah']
                        
                        reflection = {
                            'reference': f"Quran {surah}:{number}",
                            'arabic': arabic['text'],
                            'english': english['text']
                        }
                else:
                    reflection = "Remember: Verily, with hardship comes ease."
                
                # 2. Fetch random Hadith (English)
                # Using a more reliable English Hadith source
                books = ['bukhari', 'muslim']
                book = random.choice(books)
                h_url = f"https://hadith-api.vercel.app/api/books/{book}?range=1-100"
                h_res = requests.get(h_url, timeout=5)
                if h_res.ok:
                    h_json = h_res.json()
                    if h_json.get('data') and h_json['data'].get('hadiths'):
                        selected = random.choice(h_json['data']['hadiths'])
                        hadith = {
                            'text': selected.get('english', 'Description not available'),
                            'source': f"{book.capitalize()} - Hadith {selected.get('hadithNumber', '?')}"
                        }
            except Exception as q_e:
                print(f"Reflection Fetch Error: {q_e}")
                reflection = "Remember to take a moment for Dhikr today."

        except Exception as e:
            print(f"Hijri Error: {e}")
            events.append(f"Could not calculate Islamic date.")
            reflection = "Error fetching reflection."
        
        # 1. Scrape OnThisDay.com for "Significance"
        url = f"https://www.onthisday.com/day/{month_name}/{day_num}"
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            # print(f"Fetching: {url}") # Reduced interaction logging
            resp = requests.get(url, headers=headers, timeout=10)
            if resp.status_code == 200:
                soup = BeautifulSoup(resp.content, 'html.parser')
                event_items = soup.select('li.event')
                for item in event_items[:5]: 
                    events.append(item.get_text(strip=True))
                
                if not soup.select('li.event'):
                     events.append("Historical events could not be fetched.")
            else:
                events.append(f"Could not connect to history source. Status: {resp.status_code}")
        except Exception as e:
            print(f"Scraping Exception: {e}")
            events.append(f"External search failed: {str(e)}")
            
        # 2. Recommended Dua (Random with Fallback)
        dua_data = None
        try:
            # Try to get from DB
            dua = Dua.query.order_by(db.func.random()).first()
            if not dua:
                # Fallback if DB is empty or fails
                dua = Dua(
                    title="Dua for Goodness", 
                    arabic_text="رَبَّنَا آتِنَا فِي الدُّنْيَا حَسَنَةً وَفِي الآخِرَةِ حَسَنَةً وَقِنَا عَذَابَ النَّارِ",
                    english_meaning="Our Lord! Give us in this world that which is good and in the Hereafter that which is good, and save us from the torment of the Fire."
                )
            
            dua_data = {
                'title': dua.title,
                'arabic': dua.arabic_text,
                'meaning': dua.english_meaning
            }
        except Exception as e:
            print(f"Dua fetch error: {e}")
            # Hard fallback
            dua_data = {
                'title': "General Dua",
                'arabic': "الْحَمْدُ لِلَّهِ",
                 'meaning': "All praise is due to Allah."
            }
            
        # 2. Day Overview (from Day model)
        day_record = Day.query.filter_by(user_id=current_user.id, date=dt.date()).first()
        day_overview = None
        if day_record:
            day_overview = {
                'intention': day_record.intention,
                'total_score': day_record.total_score,
                'energy_level': day_record.energy_level,
                'mood': day_record.mood,
                'reflection': day_record.reflection
            }
        
        # 3. Schedule Information (Logs for this date)
        s_logs = ScheduleLog.query.filter_by(user_id=current_user.id, date=dt.date()).all()
        schedule_info = []
        for s in s_logs:
            schedule_info.append({
                'title': s.routine.title if s.routine else s.task,
                'time': s.routine.start_time.strftime('%I:%M %p') if s.routine else (s.time or '-'),
                'status': s.status,
                'location': s.routine.location if s.routine and s.routine.location else None
            })

        return jsonify({
            'date': dt.strftime('%B %d, %Y'),
            'significance': events,
            'reflection': reflection, 
            'hadith': hadith,
            'dua': dua_data,
            'day_overview': day_overview,
            'schedule': schedule_info
        })
        
    except Exception as e:
        print(f"Error in get_day_details for {date_str}: {str(e)}")
        # If it's a ValueError, it's likely the date format
        if isinstance(e, ValueError):
            return jsonify({'error': 'Invalid date format'}), 400
        return jsonify({'error': f'Failed to load details: {str(e)}'}), 500

def h_day_suffix(day):
    if 11 <= day <= 13: suffix = 'th'
    else:
        suffix = {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    return f"{day}{suffix}"

@app.route('/analytics')
@login_required
def analytics_view():
    return render_template('analytics.html')

@app.route('/api/analytics_data')
@login_required
def analytics_data():
    days = request.args.get('days', 30, type=int)
    end_date = get_today()
    start_date = end_date - timedelta(days=days-1)
    
    # 1. Initialize data structure
    date_labels = []
    total_scores = []
    habit_scores = []
    prayer_scores = []
    
    # Efficient querying:
    # Fetch all logs in range
    h_logs = HabitLog.query.join(Habit).filter(Habit.user_id == current_user.id, HabitLog.date >= start_date, HabitLog.date <= end_date, HabitLog.status == True).all()
    p_logs = PrayerLog.query.filter_by(user_id=current_user.id).filter(PrayerLog.date >= start_date, PrayerLog.date <= end_date).all()
    s_logs = ScheduleLog.query.filter_by(user_id=current_user.id, status=True).filter(ScheduleLog.date >= start_date, ScheduleLog.date <= end_date).all()
    
    # Process into dictionary by date
    from collections import defaultdict
    data_map = defaultdict(lambda: {'habit': 0, 'prayer': 0, 'schedule': 0})
    
    for log in h_logs:
        points = getattr(log.habit, 'points', 10)
        data_map[log.date]['habit'] += points
        
    for log in p_logs:
        data_map[log.date]['prayer'] += log.spiritual_score

    for log in s_logs:
        data_map[log.date]['schedule'] += log.points

    # Flatten for Charts
    current = start_date
    while current <= end_date:
        date_labels.append(current.strftime('%b %d'))
        h_score = data_map[current]['habit']
        p_score = data_map[current]['prayer']
        s_score = data_map[current]['schedule']
        
        habit_scores.append(h_score)
        prayer_scores.append(p_score)
        # We can add a separate array or just merge schedule into habits for simplicity in current charts
        # But 'total_scores' MUST include it.
        total_scores.append(h_score + p_score + s_score)
        current += timedelta(days=1)
        
    return jsonify({
        'labels': date_labels,
        'habit_scores': habit_scores,
        'prayer_scores': prayer_scores,
        'total_scores': total_scores,
        'summary': {
            'total_all_time': sum(total_scores), # rough approx for range
            'avg_daily': int(sum(total_scores) / days) if days else 0,
            'best_day': max(total_scores) if total_scores else 0
        }
    })
@app.route('/schedule/upload', methods=['GET', 'POST'])
@login_required
def upload_schedule():
    if request.method == 'POST':
        manual_text = request.form.get('manual_text')
        text_content = ""
        
        # Case 1: Manual Text Input
        if manual_text and manual_text.strip():
            text_content = manual_text
            
        # Case 2: File Upload
        elif 'file' in request.files:
            file = request.files['file']
            if file and file.filename != '':
                filename = secure_filename(file.filename)
                upload_folder = os.path.join('static', 'uploads')
                os.makedirs(upload_folder, exist_ok=True)
                filepath = os.path.join(upload_folder, filename)
                file.save(filepath)
                
                ext = os.path.splitext(filename)[1].lower()
                
                try:
                    # 1. Text Files
                    if ext == '.txt':
                        with open(filepath, 'r', encoding='utf-8') as f:
                            text_content = f.read()
                            
                    # 2. Word Documents
                    elif ext in ['.doc', '.docx']:
                        if not docx:
                             flash("Word document processing not available.", "danger")
                             return redirect(request.url)
                        doc = docx.Document(filepath)
                        full_text = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text)
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                                if row_text:
                                    full_text.append(" | ".join(row_text))
                        text_content = "\n".join(full_text)
                        
                    # 3. Excel Files
                    elif ext in ['.xls', '.xlsx']:
                        if not openpyxl:
                             flash("Excel processing not available.", "danger")
                             return redirect(request.url)
                        wb = openpyxl.load_workbook(filepath, data_only=True)
                        full_text = []
                        for sheet in wb.sheetnames:
                            ws = wb[sheet]
                            full_text.append(f"--- Sheet: {sheet} ---")
                            for row in ws.iter_rows(values_only=True):
                                row_data = [str(cell) for cell in row if cell is not None]
                                if row_data:
                                    full_text.append(" | ".join(row_data))
                        text_content = "\n".join(full_text)
                        
                    # 4. Images (OCR)
                    elif ext in ['.jpg', '.jpeg', '.png']:
                        if not pytesseract:
                             flash("Tesseract OCR not installed.", "danger")
                             return redirect(request.url)
                        text_content = pytesseract.image_to_string(Image.open(filepath))
                    
                    else:
                        flash("Unsupported file type.", "warning")
                        return redirect(request.url)
                except Exception as e:
                    flash(f"Processing Error: {str(e)}", 'danger')
                    return redirect(request.url)

        if not text_content.strip():
            flash("Please upload a file or paste your schedule text.", "warning")
            return redirect(request.url)

        # Parse and show results
        parsed_items = parse_schedule_items(text_content)
        flash(f"Processing Complete! Found {len(parsed_items)} potential schedule items.", 'success')
        return render_template('upload_result.html', parsed_items=parsed_items, raw_text=text_content)
            
    return render_template('upload_schedule.html')

def parse_schedule_items(text):
    """
    Advanced parser for schedule tables (SL, Course, Section, Day, Time Range, Room).
    Returns list of dicts: {'title', 'day', 'start', 'end'}
    """
    items = []
    lines = text.split('\n')
    days_list = ['Saturday', 'Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    
    # Improved time range pattern (e.g. 09:35am – 10:35am or 09:35-10:35)
    # Group 1: Start Time, Group 2: End Time
    range_pattern = re.compile(r'(\d{1,2}[:.]\d{2}\s*(?:[APap][Mm])?)\s*[-–]\s*(\d{1,2}[:.]\d{2}\s*(?:[APap][Mm])?)')
    single_time_pattern = re.compile(r'(\d{1,2}[:.]\d{2}\s*(?:[APap][Mm])?)|(\d{1,2}\s*[APap][Mm])')
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 5: continue
        
        # 1. Detect Day
        found_day = None
        for day in days_list:
            if day.lower() in line.lower():
                found_day = day
                break
        
        if not found_day: continue # Likely a header or noise
        
        # 2. Detect Times
        start_time = ""
        end_time = ""
        
        range_match = range_pattern.search(line)
        if range_match:
            start_time = range_match.group(1).replace('.', ':')
            end_time = range_match.group(2).replace('.', ':')
        else:
            # Try to find two separate times if range symbol is weird
            times_found = single_time_pattern.findall(line)
            # findall returns tuples if there are multiple groups, let's flatten
            times_found = [t[0] or t[1] for t in times_found if any(t)]
            if len(times_found) >= 2:
                start_time = times_found[0].replace('.', ':')
                end_time = times_found[1].replace('.', ':')
            elif len(times_found) == 1:
                start_time = times_found[0].replace('.', ':')
                # Estimate end time if missing? (+1 hour)
                end_time = start_time 

        # 3. Detect Title (Course Code)
        # Pattern: Usually something like CSC 197 or STA 240
        course_match = re.search(r'[A-Z]{2,4}\s*\d{3}', line)
        if course_match:
            title = course_match.group(0)
        else:
            # Fallback: take part of the line before the day
            title = line.split(found_day)[0].strip(' 0123456789\t-|.')
            if not title: title = "Routine Item"
            
        if found_day and start_time:
            items.append({
                'title': title,
                'day': found_day,
                'start': start_time,
                'end': end_time
            })
    
    return items

@app.route('/schedule/import/confirm', methods=['POST'])
@login_required
def import_schedule_confirm():
    schedule_name = request.form.get('schedule_name', 'Imported Schedule')
    titles = request.form.getlist('titles')
    days = request.form.getlist('days')
    starts = request.form.getlist('starts')
    ends = request.form.getlist('ends')
    
    if not titles:
        flash("No items selected to import.", "warning")
        return redirect(url_for('upload_schedule'))

    # 1. Create the new Schedule
    new_schedule = Schedule(name=schedule_name, user_id=current_user.id, is_active=True)
    
    # Deactivate existing schedules for this user
    Schedule.query.filter_by(user_id=current_user.id, is_active=True).update({'is_active': False})
    
    db.session.add(new_schedule)
    db.session.commit() # Commit to get schedule.id

    # 2. Add Routine Items
    count = 0
    for i in range(len(titles)):
        title = titles[i]
        day = days[i]
        start_str = starts[i]
        end_str = ends[i]
        
        if title and day and start_str:
            try:
                # Helper to convert "09:35am" to time object
                def parse_time_flex(t_str):
                    t_str = t_str.lower().strip()
                    if 'am' in t_str or 'pm' in t_str:
                        return datetime.strptime(t_str, '%I:%M%p').time()
                    return datetime.strptime(t_str, '%H:%M').time()

                s_time = parse_time_flex(start_str)
                e_time = parse_time_flex(end_str) if end_str else s_time
                
                item = RoutineItem(
                    schedule_id=new_schedule.id,
                    title=title,
                    day_of_week=day,
                    start_time=s_time,
                    end_time=e_time
                )
                db.session.add(item)
                count += 1
            except Exception as e:
                print(f"Skipping row {i}: {e}")
                continue
                
    db.session.commit()
    flash(f'New Schedule "{schedule_name}" created with {count} routines!', 'success')
    return redirect(url_for('schedule_view'))

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
